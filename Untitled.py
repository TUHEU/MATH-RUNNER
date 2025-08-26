from docx import Document
import random

# Function to generate hard-level math MCQs
def generate_math_mcqs(n=20):
    mcqs = []
    for i in range(n):
        q_type = random.choice(["algebra", "calculus", "geometry", "number_theory", "probability"])
        
        if q_type == "algebra":
            a, b = random.randint(2, 9), random.randint(2, 9)
            q = f"Solve for x: {a}x + {b} = 0"
            correct = f"x = {-b}/{a}"
            options = [correct, f"x = {b}/{a}", f"x = {a}/{b}", f"x = {-a}/{b}"]
        
        elif q_type == "calculus":
            n_exp = random.randint(2, 5)
            q = f"Differentiate: f(x) = x^{n_exp}"
            correct = f"{n_exp}x^{n_exp-1}"
            options = [correct, f"x^{n_exp}", f"{n_exp-1}x^{n_exp}", f"{n_exp+1}x^{n_exp-2}"]
        
        elif q_type == "geometry":
            side = random.randint(3, 12)
            q = f"Find the area of a square with side {side}."
            correct = str(side**2)
            options = [correct, str(2*side), str(side*4), str(side**3)]
        
        elif q_type == "number_theory":
            num = random.randint(10, 50)
            q = f"Find the remainder when {num}^2 is divided by 5."
            correct = str((num**2) % 5)
            options = [correct, str((num**2+1)%5), str((num**2+2)%5), str((num**2+3)%5)]
        
        elif q_type == "probability":
            total = random.randint(4, 8)
            fav = random.randint(1, total-1)
            q = f"A bag has {total} balls, {fav} of which are red. Probability of picking a red ball?"
            correct = f"{fav}/{total}"
            options = [correct, f"{total-fav}/{total}", f"1/{total}", f"{fav}/{total+1}"]
        
        random.shuffle(options)
        mcqs.append((q, options, correct))
    return mcqs

# Generate 20 math MCQs
mcqs = generate_math_mcqs(20)

# Create Word document
doc = Document()
doc.add_heading("Hard Level Math MCQs", 0)

for i, (q, opts, correct) in enumerate(mcqs, start=1):
    doc.add_paragraph(f"Q{i}. {q}")
    for j, opt in enumerate(opts):
        doc.add_paragraph(f"   {chr(65+j)}. {opt}")
    doc.add_paragraph("")

# Save document
output_docx = "/mnt/data/hard_math_mcqs.docx"
doc.save(output_docx)

output_docx
